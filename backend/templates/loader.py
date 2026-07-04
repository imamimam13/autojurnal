import json
import os
import re
import hashlib
from pathlib import Path
from typing import Optional, Any

import docx
import fitz

TEMPLATES_DIR = Path(__file__).parent

BUILTIN_TEMPLATES = {
    "standard-journal": "default-journal.json",
    "standard-textbook": "default-textbook.json",
}


def _template_path(name: str) -> Path:
    safe = re.sub(r"[^a-zA-Z0-9_-]", "_", name.lower().strip())
    return TEMPLATES_DIR / f"{safe}.json"


def list_templates() -> list[dict]:
    templates = []
    for name, filename in BUILTIN_TEMPLATES.items():
        try:
            with open(TEMPLATES_DIR / filename) as f:
                data = json.load(f)
                data["id"] = name
                data["builtin"] = True
                templates.append(data)
        except Exception:
            pass
    for path in sorted(TEMPLATES_DIR.glob("*.json")):
        if path.name in BUILTIN_TEMPLATES.values():
            continue
        try:
            with open(path) as f:
                data = json.load(f)
                data["id"] = path.stem
                data["builtin"] = False
                templates.append(data)
        except Exception:
            pass
    return templates


def load_template(template_id: str) -> Optional[dict]:
    if template_id in BUILTIN_TEMPLATES:
        path = TEMPLATES_DIR / BUILTIN_TEMPLATES[template_id]
    else:
        safe = re.sub(r"[^a-zA-Z0-9_-]", "_", template_id.lower().strip())
        path = TEMPLATES_DIR / f"{safe}.json"
    try:
        with open(path) as f:
            return json.load(f)
    except Exception:
        return None


def save_template(data: dict) -> dict:
    name = data.get("name", "untitled").strip()
    if not name:
        name = "untitled"
    safe = re.sub(r"[^a-zA-Z0-9_-]", "_", name.lower().strip())
    path = TEMPLATES_DIR / f"{safe}.json"
    with open(path, "w") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    data["id"] = path.stem
    data["builtin"] = False
    return data


def delete_template(template_id: str) -> bool:
    if template_id in BUILTIN_TEMPLATES:
        return False
    safe = re.sub(r"[^a-zA-Z0-9_-]", "_", template_id.lower().strip())
    path = TEMPLATES_DIR / f"{safe}.json"
    if path.exists():
        path.unlink()
        return True
    return False


def _extract_text_from_pdf(path: str) -> str:
    doc = fitz.open(path)
    text = "".join(page.get_text() for page in doc)
    doc.close()
    return text.strip()


def _extract_text_from_docx(path: str) -> str:
    doc = docx.Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                ct = cell.text.strip()
                if ct:
                    row_text.append(ct)
            # Deduplicate adjacent cells due to merged cells
            dedup_row = []
            for val in row_text:
                if not dedup_row or dedup_row[-1] != val:
                    dedup_row.append(val)
            if dedup_row:
                parts.append(" | ".join(dedup_row))
                
    return "\n".join(parts).strip()


def parse_upload(file_path: str) -> Optional[str]:
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".pdf":
            return _extract_text_from_pdf(file_path)
        elif ext == ".docx":
            return _extract_text_from_docx(file_path)
        elif ext in (".txt", ".md"):
            with open(file_path) as f:
                return f.read().strip()
        else:
            return None
    except Exception:
        return None
